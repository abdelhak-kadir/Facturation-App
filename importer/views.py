from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from .utils.sheet_importer import *
from .models import Livraison
from .forms import LivraisonForm
from datetime import datetime
import json
from django.template.loader import render_to_string
from django.contrib import messages
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
from django.http import HttpResponse
from django.contrib import messages
from django.shortcuts import redirect
import io
from decimal import Decimal
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json


def trigger_import_ajax(request):
    """Handle import functionality with AJAX for livraison_list page"""
    if request.method == "POST":
        try:
            import_livraisons_from_sheet()
            import_livraisons_from_sheet2()
            return JsonResponse({
                'success': True, 
                'message': '✅ Import completed successfully!'
            })
        except Exception as e:
            return JsonResponse({
                'success': False, 
                'message': f'❌ Import failed: {str(e)}'
            })
    
    return JsonResponse({'success': False, 'message': 'Invalid request method'})

def livraison_list(request):
    """Your existing livraison_list view - make sure it handles success messages"""
    
    # Handle import functionality if called via regular POST
    if request.method == "POST" and 'import_action' in request.POST:
        try:
            import_livraisons_from_sheet()
            messages.success(request, '✅ Import completed successfully!')
            import_livraisons_from_sheet2()
            messages.success(request, '✅ Import completed successfully! prrrrrr')
        except Exception as e:
            messages.error(request, f'❌ Import failed: {str(e)}')
        return redirect('livraison_list')
    
    # Your existing filtering logic here
    client_filter = request.GET.get('client', '')
    chargements_filter = request.GET.get('chargements', '')
    
    # Apply your filters to get livraisons
    livraisons = Livraison.objects.order_by('-timestamp')
    
    if client_filter:
        livraisons = livraisons.filter(client__icontains=client_filter)
    if chargements_filter:
        livraisons = livraisons.filter(chargements__icontains=chargements_filter)
    
    context = {
        'livraisons': livraisons,
        'client_filter': client_filter,
        'chargements_filter': chargements_filter,
    }
    
    return render(request, 'livraison_list.html', context)

def edit_livraison(request, pk):
    livraison = get_object_or_404(Livraison, pk=pk)
    if request.method == 'POST':
        form = LivraisonForm(request.POST, instance=livraison)
        if form.is_valid():
            form.save()
            return redirect('livraison_list')
    else:
        form = LivraisonForm(instance=livraison)
    return render(request, 'edit_livraison.html', {'form': form})


def generate_invoice(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_livraisons')
        
        if not selected_ids:
            messages.error(request, 'Aucune livraison sélectionnée.')
            return redirect('livraison_list')
        
        # Get selected livraisons
        selected_livraisons = Livraison.objects.filter(pk__in=selected_ids).order_by('date_chargement')
        
        if not selected_livraisons.exists():
            messages.error(request, 'Aucune livraison trouvée avec les IDs sélectionnés.')
            return redirect('livraison_list')
        
        # Calculate totals
        total_qte = sum(livraison.qte or 0 for livraison in selected_livraisons)
        
        # Get client info (assuming all selected livraisons are for the same client)
        main_client = selected_livraisons.first().client
        
        # Generate invoice context
        context = {
            'livraisons': selected_livraisons,
            'total_qte': total_qte,
            'client': main_client,
            'invoice_date': datetime.now().strftime('%d/%m/%Y'),
            'invoice_number': f'INV-{datetime.now().strftime("%Y%m%d")}-{selected_livraisons.first().pk}',
        }
        
        # Render invoice template
        return render(request, 'invoice.html', context)
    
    return redirect('livraison_list')
def generate_word_invoice(request):
    """
    Generate a professionally styled Word invoice document based on selected livraisons
    """
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_livraisons')
        
        if not selected_ids:
            messages.error(request, 'Aucune livraison sélectionnée.')
            return redirect('livraison_list')
        
        # Get selected livraisons
        selected_livraisons = Livraison.objects.filter(pk__in=selected_ids).order_by('date_chargement')
        
        if not selected_livraisons.exists():
            messages.error(request, 'Aucune livraison trouvée avec les IDs sélectionnés.')
            return redirect('livraison_list')
        
        # Update facturation field to "facture" for selected livraisons
        selected_livraisons.update(facturation='facture')
        
        # Get client information
        main_client = selected_livraisons.first().client or "Client Non Spécifié"
        main_ice = selected_livraisons.first().ice or ""
        
        # Generate invoice details
        invoice_date = datetime.now().strftime('%d/%m/%Y')
        invoice_number = f'{datetime.now().strftime("%Y")}.{datetime.now().strftime("%m%d")}'
        
        # Create and configure document
        doc = Document()
        _configure_document_margins(doc)
        
        # Build document sections
        _add_header_section(doc, main_client, main_ice)
        _add_invoice_title_and_details(doc, invoice_number, invoice_date)
        
        # Create and populate main invoice table
        main_table = _create_main_invoice_table(doc)
        total_ht = _populate_main_table_data(main_table, selected_livraisons)
        
        # Add spacing between tables
        doc.add_paragraph()
        
        # Create and populate totals table
        _create_totals_table(doc, total_ht)
        
        # Add amount in words
        _add_amount_in_words(doc, total_ht)
        
        # Add footer to document footer section
        _add_document_footer(doc)
        
        # Generate and return response
        return _create_response(doc, invoice_number, main_client)
    
    return redirect('livraison_list')

def _configure_document_margins(doc):
    """Configure document margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(1.2)  # Increased for footer space
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def _add_header_section(doc, main_client, main_ice):
    """Add logo/company name, client name and ICE information"""
    # Company logo or name
    logo_path = "importer/media/logo.png"
    
    try:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.5))
        doc.add_paragraph()
    except:
        # Fallback to company name if logo not found
        header_para = doc.add_paragraph()
        header_run = header_para.add_run('YOUSSRA TRANS NEGOCE')
        header_run.font.size = Pt(20)
        header_run.bold = True
        header_run.font.color.rgb = RGBColor(0, 102, 204)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Client information section
    client_para = doc.add_paragraph()
    client_run = client_para.add_run(main_client.upper())
    client_run.font.size = Pt(16)
    client_run.bold = True
    client_run.font.color.rgb = RGBColor(0, 102, 204)
    client_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # ICE information
    if main_ice:
        ice_para = doc.add_paragraph()
        ice_run = ice_para.add_run(f'ICE : {main_ice}')
        ice_run.font.size = Pt(12)
        ice_run.font.color.rgb = RGBColor(102, 102, 102)
        ice_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()


def _add_invoice_title_and_details(doc, invoice_number, invoice_date):
    """Add invoice title and details section"""
    # Invoice title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('FACTURE')
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    
    # Invoice details in a formatted table-like structure
    details_para = doc.add_paragraph()
    
    # Invoice number
    number_run = details_para.add_run(f'N° de Facture    : ')
    number_run.font.size = Pt(12)
    number_run.bold = True
    number_run.font.color.rgb = RGBColor(0, 102, 204)
    
    number_value_run = details_para.add_run(invoice_number)
    number_value_run.font.size = Pt(12)
    number_value_run.font.color.rgb = RGBColor(0, 102, 204)
    
    details_para.add_run('\n')
    
    # Invoice date
    date_run = details_para.add_run(f'Date de Facture  : ')
    date_run.font.size = Pt(12)
    date_run.bold = True
    date_run.font.color.rgb = RGBColor(0, 102, 204)
    
    date_value_run = details_para.add_run(invoice_date)
    date_value_run.font.size = Pt(12)
    date_value_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()


def _create_main_invoice_table(doc):
    """Create and style the main invoice table (without totals)"""
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    
    # Configure header row
    header_cells = table.rows[0].cells
    headers = ['Date', 'N° de BC', 'Trajet', 'Qté', 'PU HT', 'Montant HT']
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        
        # Apply header styling
        shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Style header text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Set column widths for better layout
    widths = [Inches(0.9), Inches(0.9), Inches(3.0), Inches(0.7), Inches(0.9), Inches(1.1)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width
    
    return table


def _populate_main_table_data(table, selected_livraisons):
    """Populate main table with livraison data and return total"""
    total_ht = Decimal('0.00')
    
    for idx, livraison in enumerate(selected_livraisons):
        row_cells = table.add_row().cells
        
        # Apply alternating row colors for better readability
        if idx % 2 == 0:
            for cell in row_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Populate row data
        _populate_row_data(row_cells, livraison)
        
        # Calculate amounts
        qte = livraison.qte or 0
        try:
            tarif_str = str(livraison.tarif or '0').replace(',', '').replace(' ', '')
            tarif = Decimal(tarif_str) if tarif_str.replace('.', '').isdigit() else Decimal('0')
        except:
            tarif = Decimal('0')
        
        montant = tarif * qte
        total_ht += montant
        
        # Update amount cell
        row_cells[4].text = f"{tarif:.2f}"
        row_cells[5].text = f"{montant:.2f}"
        
        # Apply cell formatting
        _format_data_cells(row_cells)
    
    return total_ht


def _populate_row_data(row_cells, livraison):
    """Populate individual row data"""
    # Date
    date_str = livraison.date_chargement.strftime('%d/%m/%Y') if livraison.date_chargement else '-'
    row_cells[0].text = date_str
    
    # BC Number
    row_cells[1].text = livraison.bon_livraison or '-'
    
    # Trajet (route description)
    trajet = f"TRANSPORT {livraison.chargements or ''} VERS {livraison.dechargement or ''}".strip()
    if trajet == "TRANSPORT  VERS":
        trajet = f"TRANSPORT {livraison.client or ''}"
    row_cells[2].text = trajet
    
    # Quantity
    qte = livraison.qte or 0
    row_cells[3].text = str(qte)


def _format_data_cells(row_cells):
    """Apply formatting to data cells"""
    for i, cell in enumerate(row_cells):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            
            # Right align numeric columns
            if i in [3, 4, 5]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _create_totals_table(doc, total_ht):
    """Create a separate table for totals aligned with PU HT and Montant HT columns"""
    # Create table with 2 columns (PU HT and Montant HT equivalent)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    
    # Set table alignment to right
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # Calculate amounts
    tva_rate = Decimal('0.20')  # 20%
    tva_amount = total_ht * tva_rate
    total_ttc = total_ht + tva_amount
    
    # Define totals data
    totals_data = [
        ("Total HT", total_ht, "4472C4"),
        ("TVA 20%", tva_amount, "4472C4"),
        ("TOTAL TTC", total_ttc, "2E5BC7")
    ]
    
    # Set column widths to match PU HT and Montant HT columns
    widths = [Inches(0.9), Inches(1.1)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width
    
    # Populate totals table
    for row_idx, (label, amount, bg_color) in enumerate(totals_data):
        row = table.rows[row_idx]
        
        # Label cell
        label_cell = row.cells[0]
        label_cell.text = label
        
        # Amount cell
        amount_cell = row.cells[1]
        amount_cell.text = f"{amount:.2f}"
        
        # Apply styling to both cells
        for cell in [label_cell, amount_cell]:
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12 if label == "TOTAL TTC" else 11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    return total_ttc


def _add_amount_in_words(doc, total_ht):
    """Add amount in words section"""
    # Calculate total TTC for amount in words
    tva_rate = Decimal('0.20')
    tva_amount = total_ht * tva_rate
    total_ttc = total_ht + tva_amount
    
    doc.add_paragraph()
    
    amount_words = f"Arrêtée La Présente Facture à la somme de {number_to_french_words(total_ttc)} Dirhams Et Toutes Taxes Comprises."
    
    words_para = doc.add_paragraph(amount_words)
    words_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for run in words_para.runs:
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.italic = True


def _add_document_footer(doc):
    """Add company information to document footer section"""
    # Access the footer
    section = doc.sections[0]
    footer = section.footer
    
    # Add a subtle border line at top of footer
    footer_border = footer.paragraphs[0]
    footer_border_run = footer_border.add_run('_' * 80)
    footer_border_run.font.color.rgb = RGBColor(200, 200, 200)
    footer_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company address
    address_para = footer.add_paragraph()
    address_run = address_para.add_run('23 Rue Bouredmg Etg 3 N° 5 Roches Noires Casablanca')
    address_run.font.size = Pt(10)
    address_run.bold = True
    address_run.font.color.rgb = RGBColor(0, 102, 204)
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company details
    details_para = footer.add_paragraph()
    details_text = ('Tél : 06.60.67.00.27-06.26.89.60.58 – RC N°591493. à Casablanca – '
                   'T.P N° 31302439 – I.F N° 53893013– ICE : 003341171000032 '
                   'Email: youssratransnegoce@gmail.com')
    
    details_run = details_para.add_run(details_text)
    details_run.font.size = Pt(9)
    details_run.font.color.rgb = RGBColor(102, 102, 102)
    details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _create_response(doc, invoice_number, main_client):
    """Create HTTP response with the generated document"""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    
    # Clean filename for better compatibility
    safe_client_name = main_client.replace(" ", "_").replace("/", "-")
    filename = f'Facture_{invoice_number}_{safe_client_name}.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

def number_to_french_words(number):
    """
    Convert a number to French words
    """
    try:
        from num2words import num2words
        # Convert to integer for cleaner output, then format
        amount_int = int(number)
        amount_cents = int((number - amount_int) * 100)
        
        words = num2words(amount_int, lang='fr').title()
        if amount_cents > 0:
            cents_words = num2words(amount_cents, lang='fr').title()
            return f"{words} Dirhams Et {cents_words} Centimes"
        else:
            return f"{words} Dirhams"
    except ImportError:
        # Fallback if num2words is not installed
        return f"{number:.2f}"


